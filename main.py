# 这是一个示例 Python 脚本。

# 按 Shift+F10 执行或将其替换为您的代码。
# 按 双击 Shift 在所有地方搜索类、文件、工具窗口、操作和设置。


import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt
import re


def set_run(run, font_size, bold, color):
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color


# 遍历文件夹下所有docx文件
def get_all_files(rootdir):
    files = []
    filelist = os.listdir(rootdir)
    for i in range(0, len(filelist)):
        path = os.path.join(rootdir, filelist[i])
        if os.path.isdir(path):
            files.extend(get_all_files(path))
        if os.path.isfile(path):
            if path.endswith(".docx"):
                files.append(path)
    # print(files)
    return files


# 查找并文件中的“答案部分”
def find_answer(file):
    pass
    # word = Dispatch('Word.Application')
    # doc = word.Documents.Open(file)
    # for i in range(1, doc.Paragraphs.Count + 1):
    #     if "答案部分" in doc.Paragraphs(i).Range.Text:
    #         return i
    # return -1


# 按装订区域中的绿色按钮以运行脚本。
if __name__ == '__main__':
    # inputFile = "D:\\1、讲义\\1、相关专业练习\\xgzs_lx1011.docx"
    # outputFile = "D:\\1、讲义\\output\\xgzs_lx1011.docx"

    outputdir = "D:\\output\\"
    for file in get_all_files("D:\\1、讲义\\"):
        # get file name with extension

        file_name = os.path.splitext(file)[0]
        # get file extension
        file_extension = os.path.splitext(file)[1]

        # 创建一个Document对象
        document = Document(file)

        paragraphs = document.paragraphs

        # 问题与答案中间有“答案部分”这4个字，先找出这一个行，作为分隔点
        pattern_daan_bufen = r"答案部分"  # pattern_daan = r"\d、\n\【正确答案】\s*(\w)"
        # 隐藏着答案“ABCDE”的那一行以“【正确答案】”开头
        pattern_daan = r"^【正确答案】\s*(\w)"
        # 每个问题的句子
        pattern_wenti = r"^(\d+)、([A-Z]|[a-z]|\d+|《|[\u4e00-\u9fa5]+)"
        # 提取题目号的表达式
        pattern_tihao = r"^(\d+)"
        pattern_timu_leixing_fenge_2 = r"^二、B"  # 读取 “二、B”
        pattern_timu_leixing_fenge_3 = r"^三、C"  # 读取 “三、C”
        pattern_tihao_leixing2 = r"^<\d+>"

        # 找到"答案部分"行前设置为0，找到后设置为1.根据此标志来确定是问题部分还是答案部分
        flag_daan_fenjie = 0

        # 题目中，“二、B” 所在行的索引
        index_timu_leixing_2 = -1
        # 题目中，“三、C” 所在行的索引
        index_timu_leixing_3 = -1

        # "答案部分"所在行的索引
        index_daan_fenge = -1

        # 答案中，“二、B” 所在行的索引
        index_daan_leixing_2 = -1
        # 答案中，“三、C” 所在行的索引
        index_daan_leixing_3 = -1

        # 遍历docx文件的所有段落，找出带有“答案部分”的那一行，行号作为问题与答案的分隔
        for i in range(len(paragraphs) - 1):
            result_timu_leixing_fenjie = re.search(pattern_timu_leixing_fenge_2, paragraphs[i].text.strip())
            if result_timu_leixing_fenjie and flag_daan_fenjie:
                index_daan_leixing_2 = i
                print(index_daan_leixing_2)
                continue
            elif result_timu_leixing_fenjie and not flag_daan_fenjie:
                index_timu_leixing_2 = i
                print(index_timu_leixing_2)

            result_timu_leixing_fenjie_3 = re.search(pattern_timu_leixing_fenge_3, paragraphs[i].text.strip())
            if result_timu_leixing_fenjie_3 and flag_daan_fenjie:
                index_daan_leixing_3 = i
                print(index_daan_leixing_3)
                continue
            elif result_timu_leixing_fenjie_3 and not flag_daan_fenjie:
                index_timu_leixing_3 = i
                print(index_timu_leixing_3)

            result = re.search(pattern_daan_bufen, paragraphs[i].text.strip())
            if result:
                index_daan_fenge = i
                flag_daan_fenjie = 1
                print(index_daan_fenge)

        index_end_loop_daan = len(paragraphs) - 1
        index_end_loop_timu = index_daan_fenge -1

        if index_daan_leixing_3 >= 0:
            index_end_loop_daan = index_daan_leixing_3-1
        if index_timu_leixing_3 >= 0:
            index_end_loop_timu = index_timu_leixing_3-1

        if index_daan_leixing_2 >= 0:
            index_end_loop_daan = index_daan_leixing_2-1
        if index_timu_leixing_2 >= 0:
            index_end_loop_timu = index_timu_leixing_2-1

        # 如果带有“答案部分”的行号找到
        if index_daan_fenge >= 0 and index_daan_fenge < len(paragraphs):
            # 从带有“答案部分”的行号+1开始遍历单选题答案
            for index in range(index_daan_fenge + 1, index_end_loop_daan):
                result = re.search(pattern_daan, paragraphs[index].text)
                if result:
                    daan = paragraphs[index].text
                    daan = daan.replace("【正确答案】", "").strip()
                    result_tihao = re.search(pattern_tihao, paragraphs[index - 1].text)
                    if result_tihao:
                        tihao = result_tihao.group(1)
                        index_wenti: int
                        for index_wenti in range(1, index_end_loop_timu):
                            result_wenti_tihao = re.search(pattern_wenti, paragraphs[index_wenti].text)
                            if result_wenti_tihao:
                                wenti_tihao = result_wenti_tihao.group(1)
                                target_wenti_index = 0
                                if wenti_tihao == tihao:
                                    match daan.strip():
                                        case str("A"):
                                            target_wenti_index = index_wenti + 1
                                        case str("B"):
                                            target_wenti_index = index_wenti + 2
                                        case str("C"):
                                            target_wenti_index = index_wenti + 3
                                        case str("D"):
                                            target_wenti_index = index_wenti + 4
                                        case str("E"):
                                            target_wenti_index = index_wenti + 5

                                    # 设置paragraphs[target_wenti_index]的文字为高亮显示
                                    run = paragraphs[target_wenti_index].runs[0]
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    break

            daan_list = []
            tihao = 0
            if index_daan_leixing_3>=0:
                index_end_loop = index_daan_leixing_3
            else:
                index_end_loop = len(paragraphs)-1

            for index in range(index_daan_leixing_2+1, index_end_loop):
                result = re.search(pattern_tihao, paragraphs[index].text)
                if result:
                    # index_list.append(index)
                    tihao = result.group(1).strip()
                    daan_list = []
                    for i in range(index + 1, index_end_loop):
                        result_1 = re.search(pattern_daan, paragraphs[i].text)
                        result_2 = re.search(pattern_tihao, paragraphs[i].text)
                        if result_1:
                            daan_list.append(result_1.group(1).strip())
                        elif result_2 or i == index_end_loop - 1:

                            for index_wenti in range(index_timu_leixing_2, index_end_loop - 1):
                                result_3 = re.search(pattern_tihao, paragraphs[index_wenti].text)
                                if result_3:
                                    wenti_tihao = result_3.group(1).strip()
                                    if wenti_tihao == tihao:

                                        for j in range(index_wenti+1, index_end_loop):
                                            result_4 = re.search(pattern_tihao_leixing2, paragraphs[j].text)
                                            result_5 = re.search(pattern_wenti, paragraphs[j].text)
                                            if result_5:
                                                break
                                            if not result_5 and result_4:
                                                if len(daan_list) >= 1:
                                                    daan = daan_list.pop(0)
                                                else:
                                                    daan = ""
                                                    break
                                                run = paragraphs[j + 1].runs[0]
                                                font_size = run.font.size  # 切片，得到一个列表rerun
                                                bold = run.bold
                                                color = run.font.color.rgb

                                                paragraphs[j + 1].text = ""  # 对run的文本清空处理

                                                run = paragraphs[j + 1].add_run(daan)

                                                set_run(run, font_size, bold, color)
                                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                                        break
                            break

            if index_daan_leixing_3 >= 0:
                daan_list = []
                tihao = 0

                for index in range(index_daan_leixing_3, len(paragraphs)):
                    result = re.search(pattern_tihao, paragraphs[index].text)
                    if result:
                        # index_list.append(index)
                        tihao = result.group(1).strip()
                        daan_list = []
                        for i in range(index + 1, len(paragraphs)):
                            result_1 = re.search(pattern_daan, paragraphs[i].text)
                            result_2 = re.search(pattern_tihao, paragraphs[i].text)
                            if result_1:
                                daan_list.append(result_1.group(1).strip())
                            elif result_2 or i == len(paragraphs) - 1:
                                # index = i
                                for index_wenti in range(index_timu_leixing_3+1, index_daan_fenge - 1):
                                    result_3 = re.search(pattern_tihao, paragraphs[index_wenti].text)
                                    if result_3:
                                        wenti_tihao = result_3.group(1).strip()
                                        if wenti_tihao == tihao:

                                            for j in range(index_wenti + 1, index_daan_fenge):
                                                result_4 = re.search(pattern_tihao_leixing2, paragraphs[j].text)
                                                result_5 = re.search(pattern_wenti, paragraphs[j].text)
                                                if result_5:
                                                    break
                                                if not result_5 and result_4:
                                                    if len(daan_list) >= 1:
                                                        daan = daan_list.pop(0)
                                                    else:
                                                        daan = ""
                                                        break

                                                    target_wenti_index = 0
                                                    match daan.strip():
                                                        case str("A"):
                                                            target_wenti_index = j + 1
                                                        case str("B"):
                                                            target_wenti_index = j + 2
                                                        case str("C"):
                                                            target_wenti_index = j + 3
                                                        case str("D"):
                                                            target_wenti_index = j + 4
                                                        case str("E"):
                                                            target_wenti_index = j + 5

                                                    # 设置paragraphs[target_wenti_index]的文字为高亮显示
                                                    run = paragraphs[target_wenti_index].runs[0]
                                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                                            break
                                break

        document.save(os.path.join(outputdir, file_name + "_output" + file_extension))

    # for file in get_all_files("D:\\1、讲义\\"):
    # save_as_docx(file)

    # word = Dispatch('Word.Application')
    # doc = word.Documents.Open(file)
    #
    # doc.SaveAs(file + "x", 12)
    # doc.Close()
    # word.Quit()
    # os.remove(file)
